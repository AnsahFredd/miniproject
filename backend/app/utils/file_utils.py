from typing import Union
from fastapi import UploadFile
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
import os


# --- Async helper for FastAPI file upload ---
async def parse_file_content(file: UploadFile) -> Union[str, None]:
    filename = file.filename.lower()
    ext = os.path.splitext(filename)[-1]

    content = await file.read()
    if not content:
        raise ValueError(f"Uploaded file {file.filename} is empty.")

    if ext == ".pdf":
        return extract_from_pdf_bytes(content)
    elif ext == ".docx":
        return extract_from_docx_bytes(content)
    elif ext == ".txt":
        return content.decode("utf-8", errors="ignore").strip()
    else:
        raise ValueError(f"Unsupported file type: {file.filename}")


# --- For use with file paths (sync) ---
def extract_text_from_file(file_path: str) -> Union[str, None]:
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".pdf":
        return extract_from_pdf(file_path)
    elif ext == ".docx":
        return extract_from_docx(file_path)
    elif ext == ".txt":
        return extract_from_txt(file_path)
    else:
        return None


# --- PDF, DOCX, TXT path-based extractors ---
def extract_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    return "\n".join([page.extract_text() or "" for page in reader.pages])


def extract_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


# --- Byte-based extractors (for UploadFile) ---
def extract_from_pdf_bytes(file_bytes: bytes) -> str:
    if not file_bytes:
        raise ValueError("Uploaded PDF is empty.")
    try:
        reader = PdfReader(BytesIO(file_bytes))
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")


def extract_from_docx_bytes(file_bytes: bytes) -> str:
    if not file_bytes:
        raise ValueError("Uploaded DOCX is empty.")
    try:
        doc = Document(BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")
